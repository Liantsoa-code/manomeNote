import re
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def set_cell_background(cell, fill_color):
    """Applique une couleur d'arrière-plan à une cellule (format RRGGBB)"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_tests_doc(content, output_file="Cahier_de_Tests.docx"):
    doc = Document()
    
    # Style de titre principal
    title = doc.add_heading('Cahier des Tests Fonctionnels', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Séparation du texte par sections (###)
    sections = re.split(r'###\s+', content)
    
    for section in sections:
        if not section.strip():
            continue
            
        lines = section.strip().split('\n')
        section_title = lines[0]
        doc.add_heading(section_title, level=1)
        
        # Extraction du tableau
        table_data = []
        for line in lines:
            if line.startswith('|') and '---' not in line:
                # Nettoyage des colonnes et suppression des balises <br>
                cols = [c.strip().replace('<br>', '\n').replace('**', '') for c in line.split('|')[1:-1]]
                if cols:
                    table_data.append(cols)
        
        if table_data:
            # Création du tableau dans Word
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            table.style = 'Table Grid'
            
            for i, row_data in enumerate(table_data):
                row_cells = table.rows[i].cells
                for j, val in enumerate(row_data):
                    row_cells[j].text = val
                    
                    # Style pour l'en-tête (Ligne 0)
                    if i == 0:
                        set_cell_background(row_cells[j], "2F547E") # Bleu Foncé
                        run = row_cells[j].paragraphs[0].runs[0]
                        run.font.color.rgb = RGBColor(255, 255, 255) # Texte Blanc
                        run.bold = True
                    
                    # Style pour la colonne Statut (✅ Passé)
                    if "✅" in val:
                        for run in row_cells[j].paragraphs[0].runs:
                            run.font.color.rgb = RGBColor(0, 128, 0) # Vert
        
        doc.add_paragraph("\n") # Espace entre sections

    doc.save(output_file)
    print(f"✅ Document sauvegardé sous : {output_file}")

# Lecture du fichier Markdown
if __name__ == "__main__":
    import os
    file_path = "FONCTIONNALITES_ET_TESTS.md"
    
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        create_tests_doc(content, "Cahier_de_Tests_Forage.docx")
    else:
        print(f"Erreur : Le fichier {file_path} est introuvable.")
